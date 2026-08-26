from __future__ import annotations

from io import BytesIO
from pathlib import Path
from typing import Any

TEXT_EXTENSIONS = {".txt", ".md", ".csv", ".json", ".log", ".yaml", ".yml"}
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".webp", ".tif", ".tiff", ".bmp"}


def _try_ocr_image(raw: bytes, lang: str = "ara+eng") -> str:
    """Optional OCR path. Requires system Tesseract + pytesseract + Pillow."""
    try:
        from PIL import Image
        import pytesseract
        image = Image.open(BytesIO(raw))
        return pytesseract.image_to_string(image, lang=lang) or ""
    except Exception:
        return ""


def _try_ocr_pdf(raw: bytes, lang: str = "ara+eng", max_pages: int = 25) -> str:
    """Optional scanned-PDF OCR. Requires pdf2image, poppler, pytesseract."""
    try:
        from pdf2image import convert_from_bytes
        import pytesseract
        pages = convert_from_bytes(raw, dpi=220, first_page=1, last_page=max_pages)
        texts = []
        for index, image in enumerate(pages, start=1):
            text = pytesseract.image_to_string(image, lang=lang) or ""
            if text.strip():
                texts.append(f"[page:{index}]\n{text.strip()}")
        return "\n\n".join(texts)
    except Exception:
        return ""


def extract_text_with_metadata(filename: str, raw: bytes, enable_ocr: bool = True) -> dict[str, Any]:
    suffix = Path(filename or "upload.txt").suffix.lower()
    ocr_used = False
    extraction_method = "plain-text"
    text = ""
    page_map: list[tuple[int, int, int]] = []

    if suffix in TEXT_EXTENSIONS:
        text = raw.decode("utf-8", errors="ignore")
    elif suffix == ".pdf":
        extraction_method = "pypdf"
        try:
            from pypdf import PdfReader
            reader = PdfReader(BytesIO(raw))
            parts = []
            offset = 0
            for page_no, page in enumerate(reader.pages, start=1):
                page_text = page.extract_text() or ""
                if page_text.strip():
                    segment = f"[page:{page_no}]\n{page_text.strip()}"
                    start = offset
                    parts.append(segment)
                    offset += len(segment) + 2
                    page_map.append((page_no, start, offset))
            text = "\n\n".join(parts)
        except Exception:
            text = ""
        if enable_ocr and len(text.strip()) < 80:
            ocr_text = _try_ocr_pdf(raw)
            if ocr_text.strip():
                text = ocr_text
                ocr_used = True
                extraction_method = "ocr-pdf-tesseract"
    elif suffix == ".docx":
        extraction_method = "python-docx"
        try:
            import docx
            document = docx.Document(BytesIO(raw))
            blocks = [p.text for p in document.paragraphs if p.text]
            for table in document.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cells:
                        blocks.append(" | ".join(cells))
            text = "\n".join(blocks)
        except Exception:
            text = ""
    elif suffix in {".xlsx", ".xlsm"}:
        extraction_method = "openpyxl"
        try:
            from openpyxl import load_workbook
            wb = load_workbook(BytesIO(raw), read_only=True, data_only=True)
            rows = []
            for ws in wb.worksheets:
                rows.append(f"Sheet: {ws.title}")
                for row in ws.iter_rows(values_only=True):
                    cells = [str(c) for c in row if c is not None]
                    if cells:
                        rows.append(" | ".join(cells))
            text = "\n".join(rows)
        except Exception:
            text = ""
    elif suffix in IMAGE_EXTENSIONS:
        extraction_method = "ocr-image-tesseract"
        if enable_ocr:
            text = _try_ocr_image(raw)
            ocr_used = bool(text.strip())
    else:
        text = raw.decode("utf-8", errors="ignore")

    return {
        "text": text or "",
        "ocr_used": ocr_used,
        "extraction_method": extraction_method,
        "page_map": page_map,
        "extension": suffix.lstrip("."),
    }


def extract_text_from_bytes(filename: str, raw: bytes) -> str:
    return extract_text_with_metadata(filename, raw, enable_ocr=True)["text"]
